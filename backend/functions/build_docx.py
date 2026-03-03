"""
Step 3 Function 6: buildDocx
Generates a .docx that exactly matches the NTU Biweekly Logbook template.

Reference: Mehra_Pratham_BiWeekly_Logbook_1.pdf (Version as of 2 January 2026)

Layout (every page):
  Header : NTU logo (left) | "Biweekly Logbook Entry" bold (center) | bottom rule
  Footer : "Version as of 2 January 2026" right-aligned grey

Page 1 body:
  1. Intro paragraph (bold, justified)
  2. Important Note heading + 2 bold numbered items
  3. Info table (no borders, colon separator, underlined values)
  4. Section A heading + instruction + bordered box
  5. Section B heading + 3-col table
  6. Section C heading + instruction + bordered box (bold-inline subheadings)
"""
import io
import os

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Constants ──────────────────────────────────────────────────────────────────

BORDER_COLOUR = "000000"
FOOTER_COLOUR = RGBColor(0x80, 0x80, 0x80)

LOGO_PATH = os.path.normpath(
    os.path.join(
        os.path.dirname(os.path.abspath(__file__)),
        "..", "assets", "logo_0.jpeg",
    )
)

SECTION_C_SUBHEADINGS = [
    "Key Achievements",
    "Main Challenge Faced",
    "What I Did Well",
    "Areas for Improvement",
]

_INTRO = (
    "This is a logbook entry to complete and email to your faculty supervisor as progress update on a "
    "biweekly basis (i.e. once every 2 weeks). Your submissions will be used as reference for your "
    "internship progress and will be assessed towards the overall performance of your internship journal."
)
_NOTE_1 = (
    "As consistency of submissions is assessed as part of your internship, all submissions should be made "
    "on time. Any delays should be informed and prior approval to be sought with your faculty supervisor. "
    "Late entries without prior approval will be considered invalid for assessment."
)
_NOTE_2 = (
    "Please exercise caution when disclosing sensitive corporate information and confirm that it complies "
    "with company policy. If your logbook entry includes such information, kindly obtain prior approval "
    "from your organisation\u2019s supervisor."
)

_SEC_A_INSTRUCTION = (
    "Briefly describe the objective(s) provided by the internship organisation, and the expected "
    "scope of work and deliverables expected from you to address the issues."
)
_SEC_C_INSTRUCTION = (
    "Illustrate the challenges that you are facing; describe areas that you have done well and "
    "areas that could be better managed."
)


# ── Low-level XML helpers ──────────────────────────────────────────────────────

def _run(para, text: str, *, bold=False, italic=False, underline=False,
         size=10, name="Arial") -> None:
    """Append a formatted run to a paragraph."""
    r = para.add_run(text)
    r.font.name  = name
    r.font.size  = Pt(size)
    r.bold       = bold
    r.italic     = italic
    r.font.underline = underline
    return r


def _para(doc, text="", *, bold=False, italic=False, underline=False,
          align=None, size=10, space_before=0, space_after=0,
          left_indent=None, first_line_indent=None) -> object:
    """Add a paragraph with a single run and common formatting."""
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Inches(left_indent)
    if first_line_indent is not None:
        p.paragraph_format.first_line_indent = Inches(first_line_indent)
    if text:
        _run(p, text, bold=bold, italic=italic, underline=underline, size=size)
    return p


def _border_all(table, sz="6"):
    """Apply single black border to every edge of a table (outer + inner)."""
    tbl  = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"),   "single")
        tag.set(qn("w:sz"),    sz)
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), BORDER_COLOUR)
        tblBorders.append(tag)
    tblPr.append(tblBorders)


def _border_none(table):
    """Remove all borders from a table."""
    tbl   = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "none")
        tblBorders.append(tag)
    tblPr.append(tblBorders)


def _cell_padding(cell, top=90, bottom=90, left=115, right=115):
    """Set internal cell padding in twips (1/1440 inch)."""
    tcPr  = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement("w:tcMar")
    for side, val in (("top", top), ("bottom", bottom),
                      ("left", left), ("right", right)):
        m = OxmlElement(f"w:{side}")
        m.set(qn("w:w"),    str(val))
        m.set(qn("w:type"), "dxa")
        tcMar.append(m)
    tcPr.append(tcMar)


def _shade_cell(cell, fill_hex: str):
    """Fill a table cell with a background colour."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  fill_hex)
    tcPr.append(shd)


def _para_bottom_border(para):
    """Add a solid bottom border line to a paragraph (header separator)."""
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot  = OxmlElement("w:bottom")
    bot.set(qn("w:val"),   "single")
    bot.set(qn("w:sz"),    "6")
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), "000000")
    pBdr.append(bot)
    pPr.append(pBdr)


def _center_tab(para, inches: float):
    """Add a centered tab stop to a paragraph."""
    pPr  = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab  = OxmlElement("w:tab")
    tab.set(qn("w:val"), "center")
    tab.set(qn("w:pos"), str(int(inches * 1440)))
    tabs.append(tab)
    pPr.append(tabs)


# ── Document section builders ──────────────────────────────────────────────────

def _build_header(section, logo_path: str):
    """
    Header (repeats on every page):
      [NTU logo]  <TAB>  Biweekly Logbook Entry (bold, centered)
      ─────────────────────────────────────────── (bottom rule)
    """
    header      = section.header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Center tab at the midpoint of the text area
    # A4 (8.27") − 2×1" margins = 6.27" → center at 3.135"
    _center_tab(header_para, inches=3.14)

    # NTU logo (inline image, left side)
    if os.path.exists(logo_path):
        logo_run = header_para.add_run()
        logo_run.add_picture(logo_path, height=Inches(0.55))

    # Tab to center, then bold title
    header_para.add_run("\t")
    _run(header_para, "Biweekly Logbook Entry", bold=True, size=11)

    # Horizontal rule below header
    _para_bottom_border(header_para)


def _build_footer(section):
    """Footer: 'Version as of 2 January 2026' right-aligned in grey."""
    footer      = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = footer_para.add_run("Version as of 2 January 2026")
    r.font.name       = "Arial"
    r.font.size       = Pt(8)
    r.font.color.rgb  = FOOTER_COLOUR


def _build_intro(doc: Document):
    """
    Bold justified intro paragraph + Important Note with 2 bold numbered items.
    This is static NTU boilerplate that appears on every logbook.
    """
    # Intro paragraph
    p = _para(doc, _INTRO, bold=True,
              align=WD_ALIGN_PARAGRAPH.JUSTIFY,
              space_before=4, space_after=8)

    # "Important Note:" heading
    _para(doc, "Important Note:", bold=True, space_before=4, space_after=2)

    # Numbered items (bold, hanging indent)
    for number, text in (("1.", _NOTE_1), ("2.", _NOTE_2)):
        np = _para(doc, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                   left_indent=0.4, first_line_indent=-0.25,
                   space_before=0, space_after=4)
        _run(np, number + "  ", bold=True)
        _run(np, text,          bold=True)


def _build_info_table(doc: Document, metadata: dict):
    """
    3-column info table (no borders):
      [Student Name]  [:]  [underlined value]

    Exactly 6 rows as per the NTU template.
    """
    fields = [
        ("Student Name",         metadata.get("student_name", "")),
        ("Matriculation Number", metadata.get("matric_number", "")),
        ("Company Name",         metadata.get("company", "")),
        ("Company Supervisor",   metadata.get("supervisor", "")),
        ("Logbook Entry Period",
         f"{metadata.get('period_start', '')} to {metadata.get('period_end', '')}"),
        ("Date of Submission",   metadata.get("submission_date", "")),
    ]

    tbl = doc.add_table(rows=0, cols=3)
    _border_none(tbl)

    for label, value in fields:
        row = tbl.add_row()

        # Label column
        row.cells[0].width = Inches(2.1)
        _run(row.cells[0].paragraphs[0], label)

        # Colon column
        row.cells[1].width = Inches(0.25)
        cp = row.cells[1].paragraphs[0]
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(cp, ":")

        # Value column (underlined)
        _run(row.cells[2].paragraphs[0], value, underline=True)

    _para(doc)  # spacer after info table


def _section_heading(doc: Document, letter: str, title: str):
    """
    Bold section heading:  "A.   State the Internship Objective..."
    """
    h = _para(doc, space_before=8, space_after=1)
    _run(h, f"{letter}.   {title}", bold=True)
    return h


def _section_instruction(doc: Document, text: str):
    """Indented regular instruction text under the section heading."""
    _para(doc, text,
          align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          left_indent=0.4,
          space_before=0, space_after=4)


def _build_section_a(doc: Document, section_a: str):
    """Section A: heading + instruction + bordered prose box."""
    _section_heading(doc, "A",
                     "State the Internship Objective, Scope, and Problem Statement.")
    _section_instruction(doc, _SEC_A_INSTRUCTION)

    tbl  = doc.add_table(rows=1, cols=1)
    _border_all(tbl)
    cell = tbl.cell(0, 0)
    _cell_padding(cell, top=100, bottom=100, left=130, right=130)

    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    _run(p, section_a)

    _para(doc)  # spacer


def _build_section_b(doc: Document, work_rows: list[dict]):
    """Section B: heading + 3-column work log table."""
    _section_heading(doc, "B", "Describe your work done over the past 2 weeks.")

    col_labels  = ["Work and/or Tasks Performed", "Date (From)", "Date (To)"]
    col_widths  = [Inches(4.15), Inches(1.06), Inches(1.06)]

    tbl = doc.add_table(rows=1, cols=3)
    _border_all(tbl)

    # Header row
    hdr = tbl.rows[0].cells
    for i, (label, w) in enumerate(zip(col_labels, col_widths)):
        hdr[i].width = w
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _run(p, label, bold=True)

    # Data rows
    for rd in work_rows:
        row   = tbl.add_row()
        cells = row.cells
        cells[0].width = col_widths[0]
        r = _run(cells[0].paragraphs[0], rd["task_description"])
        if rd.get("is_leave"):
            r.italic = True
        for j, key in enumerate(("date_from", "date_to"), start=1):
            dp = cells[j].paragraphs[0]
            dp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _run(dp, rd.get(key, ""))

    _para(doc)  # spacer


def _parse_section_c(text: str) -> list[dict]:
    """
    Parse Section C into paragraph segments for rendering.

    Claude outputs subheadings on their own line followed by the content.
    We merge them inline to match the PDF format:
      {'type': 'mixed',   'bold_text': 'Key Achievements:', 'rest': ' The past two weeks...'}
      {'type': 'regular', 'text': '...'}
    """
    lines = [ln.strip() for ln in text.strip().split("\n")]
    out   = []
    i     = 0

    while i < len(lines):
        line = lines[i]
        if not line:
            i += 1
            continue

        matched = next(
            (sh for sh in SECTION_C_SUBHEADINGS if line.lower().startswith(sh.lower())),
            None,
        )

        if matched:
            after = line[len(matched):].lstrip(":").strip()

            if after:
                # "Key Achievements: content on same line"
                out.append({"type": "mixed",
                            "bold_text": matched + ":",
                            "rest": " " + after})
                i += 1
            else:
                # Subheading alone — look ahead for next non-empty, non-subheading line
                j = i + 1
                while j < len(lines) and not lines[j].strip():
                    j += 1

                is_next_subheading = j < len(lines) and any(
                    lines[j].lower().startswith(sh.lower())
                    for sh in SECTION_C_SUBHEADINGS
                )

                if j < len(lines) and not is_next_subheading:
                    out.append({"type": "mixed",
                                "bold_text": matched + ":",
                                "rest": " " + lines[j].strip()})
                    i = j + 1
                else:
                    out.append({"type": "bold", "text": matched + ":"})
                    i += 1
        else:
            out.append({"type": "regular", "text": line})
            i += 1

    return out


def _build_section_c(doc: Document, section_c: str):
    """Section C: heading + instruction + bordered box with bold-inline subheadings."""
    _section_heading(
        doc, "C",
        "Describe and reflect on your work/task outcome/learning over the last 2 weeks.",
    )
    _section_instruction(doc, _SEC_C_INSTRUCTION)

    tbl  = doc.add_table(rows=1, cols=1)
    _border_all(tbl)
    cell = tbl.cell(0, 0)
    _cell_padding(cell, top=100, bottom=100, left=130, right=130)

    segments  = _parse_section_c(section_c)
    first     = True

    for seg in segments:
        if first:
            para  = cell.paragraphs[0]
            first = False
        else:
            para = cell.add_paragraph()

        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        if seg["type"] == "mixed":
            para.paragraph_format.space_before = Pt(6)
            _run(para, seg["bold_text"], bold=True)
            _run(para, seg["rest"])
        elif seg["type"] == "bold":
            para.paragraph_format.space_before = Pt(6)
            _run(para, seg["text"], bold=True)
        else:
            _run(para, seg["text"])


# ── Public entry point ─────────────────────────────────────────────────────────

def buildDocx(
    section_a: str,
    work_rows: list[dict],
    section_c: str,
    metadata: dict,
) -> bytes:
    """
    Build a .docx matching the NTU Biweekly Logbook template.

    Args:
        section_a:  Generated Section A prose text
        work_rows:  [{task_description, date_from, date_to, is_leave}]
        section_c:  Generated Section C text (with standard subheadings)
        metadata:   Student metadata dict

    Returns:
        bytes: Complete .docx file content
    """
    doc = Document()

    # ── Page setup: A4, 1" margins ──────────────────────────────────────────
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width  = Cm(21.0)
    for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
        setattr(sec, attr, Inches(1.0))

    # ── Document default font ───────────────────────────────────────────────
    style = doc.styles["Normal"]
    style.font.name = "Arial"
    style.font.size = Pt(10)

    # ── Header & footer (repeat on every page) ──────────────────────────────
    _build_header(sec, LOGO_PATH)
    _build_footer(sec)

    # ── Body ────────────────────────────────────────────────────────────────
    _build_intro(doc)
    _build_info_table(doc, metadata)
    _build_section_a(doc, section_a)
    _build_section_b(doc, work_rows)
    _build_section_c(doc, section_c)

    # ── Serialize ───────────────────────────────────────────────────────────
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()
